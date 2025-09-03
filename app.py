# app.py
# Apex AI Suite — Purple & Black Streamlit AI App
# Fixes: safe secrets handling, no duplicate imports, Python 3.8+ typing, lazy OpenAI client,
# tighter error handling, stable defaults, and minor polish.

import os
import base64
import hashlib
import sqlite3
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Tuple, Dict, Any

import numpy as np
import pandas as pd
from sklearn.ensemble import IsolationForest

import streamlit as st
from openai import OpenAI

# For exports
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib.units import cm

# -------------------------
# CONFIG & THEME
# -------------------------
st.set_page_config(
    page_title="Apex AI Suite",
    page_icon="🟣",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Purple (#6A0DAD) + Black theme CSS
st.markdown("""
<style>
:root {
  --apx-purple: #6A0DAD;
  --apx-bg: #0B0B0F;
  --apx-surface: #141420;
  --apx-text: #EAEAF2;
  --apx-muted: #B7B8C7;
}
html, body, .block-container { background: var(--apx-bg); color: var(--apx-text); }
[data-testid="stSidebar"] { background: var(--apx-surface)!important; }
h1,h2,h3 { color: var(--apx-text); }
.apx-card {
  background: var(--apx-surface); border: 1px solid rgba(255,255,255,0.06);
  border-radius: 18px; padding: 16px 18px; box-shadow: 0 8px 20px rgba(0,0,0,0.35);
}
.apx-hero {
  background: linear-gradient(90deg, rgba(106,13,173,0.2), rgba(106,13,173,0.05));
  border: 1px solid rgba(106,13,173,0.3);
  border-radius: 18px; padding: 14px 16px; margin-bottom: 10px;
}
.apx-pill { background: rgba(106,13,173,0.25); padding: 2px 8px; border-radius: 999px; font-size: 12px; }
.stButton>button, .stDownloadButton>button {
  background: var(--apx-purple)!important; color: white!important; border-radius: 10px!important; border: none!important;
}
input, textarea, .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div,
.stMultiSelect>div>div, .stFileUploader {
  background: #0F0F16!important; color: var(--apx-text)!important; border-radius: 10px!important; border: 1px solid #1f1f2a!important;
}
</style>
""", unsafe_allow_html=True)

# -------------------------
# DATABASE SETUP (persistent users with hashed passwords)
# -------------------------
DB_PATH = Path("users.db")
conn = sqlite3.connect(DB_PATH, check_same_thread=False)
cur = conn.cursor()
cur.execute('''CREATE TABLE IF NOT EXISTS users (email TEXT PRIMARY KEY, password_hash TEXT)''')
conn.commit()

def hash_password(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def add_user(email: str, pw: str) -> bool:
    pw_hash = hash_password(pw)
    try:
        cur.execute("INSERT INTO users (email, password_hash) VALUES (?, ?)", (email, pw_hash))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False

def verify_user(email: str, pw: str) -> bool:
    cur.execute("SELECT password_hash FROM users WHERE email = ?", (email,))
    row = cur.fetchone()
    if row:
        return hash_password(pw) == row[0]
    return False

# -------------------------
# SECRETS / ENV (safe)
# -------------------------
def get_secret(key: str, default: Optional[str] = None) -> Optional[str]:
    try:
        # Accessing st.secrets can raise if no secrets.toml exists
        return st.secrets[key]  # type: ignore[index]
    except Exception:
        return os.getenv(key, default)

OPENAI_API_KEY: Optional[str] = get_secret("OPENAI_API_KEY")

# -------------------------
# OPENAI SETUP
# -------------------------
PRIMARY_MODEL = "gpt-4o-mini"
IMAGE_MODEL = "dall-e-3"  # OpenAI Images API
TEMPERATURE_DEFAULT = 0.35

def get_openai_client() -> Optional[OpenAI]:
    if "openai_client" in st.session_state:
        return st.session_state.openai_client
    api_key = OPENAI_API_KEY
    if not api_key:
        st.warning("OpenAI API key not set. Add it in `.streamlit/secrets.toml` as `OPENAI_API_KEY` or set an environment variable.")
        return None
    try:
        cli = OpenAI(api_key=api_key)
        st.session_state.openai_client = cli
        return cli
    except Exception as e:
        st.error(f"OpenAI init error: {e}")
        return None

def ai_complete(
    prompt: str,
    system: str = "You are a precise, structured business and productivity expert.",
    model: str = PRIMARY_MODEL,
    temperature: float = TEMPERATURE_DEFAULT
) -> Optional[str]:
    client = get_openai_client()
    if client is None:
        return None
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            temperature=float(temperature)
        )
        content = resp.choices[0].message.content
        return content.strip() if content else ""
    except Exception as e:
        st.error(f"OpenAI completion error: {e}")
        return None

def ai_generate_logo(prompt: str, size: str = "512x512") -> Optional[bytes]:
    """Use OpenAI Images API to generate a logo; returns PNG bytes or None."""
    client = get_openai_client()
    if client is None:
        return None
    try:
        resp = client.images.generate(
            model=IMAGE_MODEL,
            prompt=prompt,
            size=size,
            n=1,
            response_format="b64_json"
        )
        b64 = resp.data[0].b64_json
        return base64.b64decode(b64) if b64 else None
    except Exception as e:
        st.info(f"Logo generation error: {e}")
        return None

# -------------------------
# EXPORTERS (DOCX/PDF)
# -------------------------
def export_to_docx(title: str, sections: List[Tuple[str, str]]) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for h, t in sections:
        doc.add_heading(h, level=2)
        for para in t.split("\n"):
            p = doc.add_paragraph(para.strip())
            for run in p.runs:
                run.font.size = Pt(11)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def export_to_pdf(title: str, sections: List[Tuple[str, str]]) -> bytes:
    bio = BytesIO()
    c = pdf_canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    y = height - margin

    def wrap(text: str, x: float, y_val: float, w: float, font: str = "Helvetica", size: int = 10) -> float:
        c.setFont(font, size)
        line_h = size + 2
        for line in text.split("\n"):
            words = line.split(" ")
            cur = ""
            for word in words:
                test = (cur + " " + word).strip()
                if c.stringWidth(test, font, size) > w:
                    c.drawString(x, y_val, cur)
                    y_val -= line_h
                    cur = word
                else:
                    cur = test
            if cur:
                c.drawString(x, y_val, cur)
                y_val -= line_h
        return y_val

    c.setFont("Helvetica-Bold", 16)
    y = wrap(title, margin, y, width - 2 * margin)
    y -= 24
    for h, t in sections:
        if y < 3 * cm:
            c.showPage()
            y = height - margin
        y = wrap(h, margin, y, width - 2 * margin, "Helvetica-Bold", 12)
        y -= 16
        y = wrap(t, margin, y, width - 2 * margin)
        y -= 10
    c.save()
    bio.seek(0)
    return bio.read()

def download_buttons(file_title: str, sections: List[Tuple[str, str]]):
    with st.container():
        col1, col2 = st.columns(2)
        docx_bytes = export_to_docx(file_title, sections)
        pdf_bytes = export_to_pdf(file_title, sections)
        col1.download_button(
            "📄 Download DOCX",
            docx_bytes,
            file_name=f"{file_title.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        col2.download_button(
            "🧾 Download PDF",
            pdf_bytes,
            file_name=f"{file_title.replace(' ', '_')}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

# -------------------------
# UI HELPERS
# -------------------------
def hero_bar():
    st.markdown("""
    <div class="apx-hero">
      <div style="display:flex;justify-content:space-between;align-items:center;">
        <div>
          <div class="apx-pill">Apex AI Suite</div>
          <h2 style="margin:6px 0 0 0;">🟣 Purple & Black AI Productivity Platform</h2>
          <div style="color:#cfcfe3;font-size:13px;margin-top:4px;">GPT-powered • Modular • DOCX/PDF exports</div>
        </div>
        <div style="font-size:12px;color:#cfcfe3;">Streamlit + OpenAI</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

def header(title: str, subtitle: str = ""):
    st.markdown(f"""
      <div class="apx-card" style="margin-bottom:10px;">
        <h3 style="margin:0 0 6px 0;">{title}</h3>
        <div style="color:#c7c7d9;font-size:13px;">{subtitle}</div>
      </div>
    """, unsafe_allow_html=True)

# Shortcuts
def tip(msg): st.info(msg)
def warn(msg): st.warning(msg)
def err(msg): st.error(msg)
def ok(msg): st.success(msg)

# -------------------------
# AUTH
# -------------------------
if "auth" not in st.session_state:
    st.session_state.auth = {"logged_in": False, "user": None}

def login_page():
    hero_bar()
    header("🔐 Login", "Access your modules")
    email = st.text_input("Email", placeholder="you@example.com")
    pw = st.text_input("Password", type="password")
    if st.button("Login", type="primary"):
        if verify_user(email, pw):
            st.session_state.auth = {"logged_in": True, "user": email}
            ok("Logged in ✅")
            st.rerun()
        else:
            err("Invalid credentials ❌")

def signup_page():
    hero_bar()
    header("📝 Sign Up", "Create a new account")
    email = st.text_input("New Email")
    pw = st.text_input("Password", type="password")
    if st.button("Create Account", type="primary"):
        if add_user(email, pw):
            ok("Account created ✅ Now log in.")
        else:
            err("User already exists")

def logout():
    st.session_state.auth = {"logged_in": False, "user": None}
    ok("You have been logged out.")
    st.rerun()

# -------------------------
# MARKETPLACE
# -------------------------
def page_marketplace():
    hero_bar()
    header("🛒 Marketplace", "Browse modules and install (placeholder)")
    tools = [
        ("🚀 Startup Copilot", "Validate & structure your idea"),
        ("📊 Business Ops Monitor", "CFO analysis & anomaly detection"),
        ("🎯 Work Productivity Coach", "Weekly plans & deep work"),
        ("💹 AI Investment Assistant", "Diversification & rebalance"),
        ("📈 AI Stock Trading Journal", "Biases & performance patterns"),
        ("📝 Ad Copy & Creative Generator", "PAS/AIDA + hooks & creatives"),
        ("📄 AI CV Optimizer", "ATS-friendly tailored CVs"),
        ("💼 Job Search Copilot", "Roles, letters, recruiter DMs"),
        ("🏡 Remote Work Companion", "Focus plan & wellness"),
        ("✍️ AI Content Writer", "Blogs, SEO, LinkedIn, Threads"),
        ("📢 AI Ad Generator", "Multi-platform campaigns"),
        ("🎨 AI Logo Creator", "Prompt → logo concepts"),
        ("🔍 AI Logo Scanner", "Upload → brand critique"),
    ]
    cols = st.columns(3)
    for i, (name, desc) in enumerate(tools):
        with cols[i % 3]:
            st.markdown(f"""
            <div class="apx-card">
              <div style="font-weight:700;margin-bottom:4px;">{name}</div>
              <div style="color:#c7c7d9;font-size:13px;margin-bottom:8px;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Install {name}", key=f"install_{i}"):
                tip(f"Installed {name} (placeholder action)")

# -------------------------
# MODULES
# -------------------------
def page_content_writer():
    hero_bar()
    header("✍️ AI Content Writer", "Blogs, SEO, LinkedIn posts, Threads")
    topic = st.text_input("Topic", placeholder="e.g., Automating SME ops with AI")
    style = st.selectbox("Style", ["Blog", "SEO Article", "LinkedIn Post", "Twitter Thread"])
    words = st.slider("Target length (words)", 150, 2000, 700)
    if st.button("Generate Content", type="primary"):
        if not topic:
            warn("Please enter a topic.")
            return
        prompt = f"Write a {words}-word {style} about: {topic}. Use clear sections, engaging tone, and actionable takeaways."
        out = ai_complete(prompt, temperature=0.6)
        if out:
            st.markdown(out)
            sections = [("Topic", topic), ("Style", style), ("AI Content", out)]
            download_buttons("AI_Content_Writer_Output", sections)
        else:
            st.markdown("—")

def page_ai_ad_generator():
    hero_bar()
    header("📢 AI Ad Generator", "Campaign variations per platform")
    product = st.text_input("Product/Service")
    audience = st.text_input("Target Audience", placeholder="e.g., busy founders in Lagos")
    platforms = st.multiselect("Platforms", ["Facebook", "Instagram", "TikTok", "LinkedIn", "X/Twitter", "YouTube"], default=["Instagram", "LinkedIn"])
    if st.button("Generate Campaign", type="primary"):
        if not product or not audience:
            warn("Add product and audience.")
            return
        prompt = f"""
        Create a multi-platform ad campaign for "{product}" to "{audience}" on {', '.join(platforms)}.
        Include: 5 hooks, 3 headline variants per platform, primary text, 5 CTAs, and creative suggestions per platform.
        """
        out = ai_complete(prompt, temperature=0.7)
        if out:
            st.markdown(out)
            sections = [("Product", product), ("Audience", audience), ("Platforms", ", ".join(platforms)), ("AI Campaign", out)]
            download_buttons("AI_Ad_Generator_Campaign", sections)
        else:
            st.markdown("—")

def page_logo_creator():
    hero_bar()
    header("🎨 AI Logo Creator", "Generate logo concepts from a prompt")
    desc = st.text_input("Describe your brand/logo idea", placeholder="Name, vibe, industry; prefer purple/black")
    size = st.selectbox("Image size", ["256x256", "512x512", "1024x1024"], index=1)
    if st.button("Generate Logo Concept", type="primary"):
        if not desc:
            warn("Please describe your brand/logo.")
            return
        prompt = f"Minimal, professional vector-style logo. {desc}. Purple and black palette."
        img_bytes = ai_generate_logo(prompt, size=size)
        if img_bytes:
            st.image(img_bytes, caption="AI Logo Concept", use_column_width=False)
            st.download_button("Download PNG", img_bytes, file_name="logo_concept.png", mime="image/png")
        else:
            st.info("Image API unavailable; showing placeholder.")
            st.image("https://placehold.co/512x512/6A0DAD/FFFFFF?text=Logo+Concept")
        sections = [("Logo Prompt", desc), ("Note", "Enable OpenAI Images API in production.")]
        download_buttons("Logo_Creator_Output", sections)

def page_logo_scanner():
    hero_bar()
    header("🔍 AI Logo Scanner", "Upload a logo for critique & suggestions")
    file = st.file_uploader("Upload logo image", type=["png", "jpg", "jpeg"])
    if file:
        st.image(file, caption="Uploaded Logo", width=220)
        if st.button("Analyze Logo", type="primary"):
            prompt = """
            Provide a professional logo critique given a hypothetical uploaded logo:
            - Style & visual identity
            - Color psychology (emphasize purple/black if relevant)
            - Simplicity & memorability
            - Legibility at small sizes
            - Brand fit & differentiation
            - 6–10 actionable improvements
            - Export-ready notes
            """
            out = ai_complete(prompt, temperature=0.35)
            if out:
                st.markdown(out)
                sections = [("Observations", "See below"), ("AI Critique", out)]
                download_buttons("Logo_Scanner_Report", sections)
            else:
                st.markdown("—")

def page_startup_copilot():
    hero_bar()
    header("🚀 Startup Copilot", "From idea to execution plan")
    idea = st.text_area("Your business idea")
    market = st.text_input("Target market")
    tone = st.selectbox("Tone", ["Concise", "Detailed", "Investor-ready"], index=1)
    depth = st.slider("Depth", 1, 3, 2)
    if st.button("Generate Plan", type="primary"):
        if not idea or not market:
            warn("Add your idea and market.")
            return
        prompt = f"""
        Business Idea: {idea}
        Target Market: {market}
        Tone: {tone}; Depth: {depth}
        Create:
        1) Lean Canvas
        2) SWOT
        3) Market insights (size, growth)
        4) 5 competitors + snapshot
        5) 90-day execution roadmap
        Format in clear markdown.
        """
        out = ai_complete(prompt, temperature=0.25)
        if out:
            st.markdown(out)
            sections = [("Idea", idea), ("Target Market", market), ("AI Plan", out)]
            download_buttons("Startup_Copilot_Report", sections)
        else:
            st.markdown("—")

def analyze_financial_csv(df: pd.DataFrame) -> Tuple[pd.DataFrame, Optional[str]]:
    num_df = df.select_dtypes(include=[np.number]).dropna()
    anomalies = pd.DataFrame()
    if not num_df.empty and len(num_df) >= 10:
        try:
            iso = IsolationForest(n_estimators=200, contamination=0.05, random_state=42)
            preds = iso.fit_predict(num_df)
            anomalies = df.loc[preds == -1].copy()
        except Exception as e:
            st.info(f"Anomaly detection error: {e}")
    prompt = f"""
    You are a CFO assistant. Analyze the uploaded financial table.
    Columns: {list(df.columns)}
    Sample:
    {df.head(30).to_csv(index=False)}
    Provide a financial health summary, key KPIs, anomalies (if any), and 10 actionable recommendations.
    """
    ai_text = ai_complete(prompt, temperature=0.25)
    return anomalies, ai_text

def page_business_ops_monitor():
    hero_bar()
    header("📊 Business Ops Monitor", "CFO-grade insights from your CSV")
    file = st.file_uploader("Upload CSV (finance/ops)", type=["csv"])
    if file:
        try:
            df = pd.read_csv(file)
        except Exception:
            file.seek(0)
            df = pd.read_csv(file, encoding="latin-1")
        st.dataframe(df, use_container_width=True, height=300)
        if st.button("Run AI Analysis", type="primary"):
            with st.spinner("Analyzing..."):
                anomalies, ai_text = analyze_financial_csv(df)
            st.subheader("AI Financial Summary")
            st.write(ai_text or "—")
            if not anomalies.empty:
                st.subheader("Potential Anomalies")
                st.dataframe(anomalies, use_container_width=True)
            else:
                tip("No strong anomalies flagged.")
            sections = [
                ("Columns", ", ".join(df.columns.astype(str))),
                ("AI Summary", ai_text or "—"),
                ("Anomalies (first 20 rows)", anomalies.head(20).to_csv(index=False) if not anomalies.empty else "None"),
            ]
            download_buttons("Business_Ops_Monitor_Report", sections)

def page_work_productivity_coach():
    hero_bar()
    header("🎯 Work Productivity Coach", "Turn chaos into a focused plan")
    goals = st.text_area("Weekly goals (comma or line-separated)")
    tasks = st.text_area("Task list (one per line; optional estimates like 'Task | 90m')")
    hours = st.slider("Focus hours/day", 1, 10, 4)
    week_start = st.date_input("Week start", datetime.today().date())
    if st.button("Generate AI Plan", type="primary"):
        prompt = f"""
        Weekly goals:
        {goals}
        Tasks:
        {tasks}
        Availability: {hours} hours/day
        Week starting: {week_start}
        Output:
        - Prioritized plan with rationale
        - Deep-work blocks (calendar-like)
        - Daily focus themes
        - Status summary template
        """
        out = ai_complete(prompt, temperature=0.25)
        if out:
            st.markdown(out)
            sections = [("Goals", goals), ("Tasks", tasks), ("AI Plan", out)]
            download_buttons("Work_Productivity_Plan", sections)
        else:
            st.markdown("—")

def parse_portfolio(text: str) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for line in text.splitlines():
        if not line.strip():
            continue
        parts = [p.strip() for p in (line.split("|") if "|" in line else line.split(","))]
        if len(parts) >= 2:
            try:
                w = float(parts[1].replace("%", "").strip())
            except ValueError:
                w = np.nan
            rows.append({"Asset": parts[0], "Weight(%)": w})
    return pd.DataFrame(rows)

def page_investment_assistant():
    hero_bar()
    header("💹 AI Investment Assistant", "Clarity on diversification and risks")
    example = "VOO | 40%\nAAPL | 15%\nQQQ | 15%\nEM ETF | 10%\nBND | 10%\nCASH | 10%"
    portfolio_text = st.text_area("Portfolio (Ticker/Asset | Weight%)", value=example, height=160)
    risk = st.select_slider("Risk preference", ["Very Low", "Low", "Moderate", "High", "Very High"], value="Moderate")
    horizon = st.selectbox("Horizon", ["<1y", "1-3y", "3-5y", "5-10y", "10y+"])
    if st.button("Analyze & Suggest Rebalance", type="primary"):
        dfp = parse_portfolio(portfolio_text)
        st.dataframe(dfp, use_container_width=True)
        total = dfp["Weight(%)"].fillna(0).sum()
        if total <= 0:
            warn("Weights could not be parsed.")
            return
        prompt = f"""
        Portfolio:
        {dfp.to_csv(index=False)}
        Total weights: {total:.1f}%
        Risk: {risk}; Horizon: {horizon}
        Provide:
        - Diversification analysis (equity/bond/intl/cash, concentration)
        - Risk commentary in plain English
        - Suggested target weights (rebalance)
        - 8 key action items
        """
        out = ai_complete(prompt, temperature=0.25)
        if out:
            st.markdown(out)
            sections = [("Portfolio", dfp.to_csv(index=False)), ("AI Suggestions", out)]
            download_buttons("Investment_Assistant_Report", sections)
        else:
            st.markdown("—")

def summarize_trades(df: pd.DataFrame) -> Dict[str, Any]:
    lower = {c.lower(): c for c in df.columns}
    candidates = ["p/l", "pnl", "profit", "return", "pl", "net"]
    pl_col = next((lower[k] for k in candidates if k in lower), df.columns[-1])
    s = df[pl_col]
    if not pd.api.types.is_numeric_dtype(s):
        s = pd.to_numeric(s, errors="coerce")
    win = (s > 0).mean() if s.notna().any() else np.nan
    avg_win = s[s > 0].mean()
    avg_loss = s[s <= 0].mean()
    expectancy = s.mean()
    return {
        "trades_n": int(len(df)),
        "win_rate": float(win) if pd.notna(win) else np.nan,
        "avg_win": float(avg_win) if pd.notna(avg_win) else np.nan,
        "avg_loss": float(avg_loss) if pd.notna(avg_loss) else np.nan,
        "expectancy": float(expectancy) if pd.notna(expectancy) else np.nan,
        "pl_col": pl_col,
    }

def page_trading_journal():
    hero_bar()
    header("📈 AI Stock Trading Journal", "Spot patterns and fix biases")
    file = st.file_uploader("Upload trading history CSV", type=["csv"])
    if file:
        try:
            df = pd.read_csv(file)
        except Exception:
            file.seek(0)
            df = pd.read_csv(file, encoding="latin-1")
        st.dataframe(df, use_container_width=True, height=320)
        m = summarize_trades(df)
        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("Trades", m["trades_n"])
        c2.metric("Win Rate", f"{m['win_rate']*100:.1f}%" if pd.notna(m["win_rate"]) else "—")
        c3.metric("Avg Win", f"{m['avg_win']:.2f}" if pd.notna(m["avg_win"]) else "—")
        c4.metric("Avg Loss", f"{m['avg_loss']:.2f}" if pd.notna(m["avg_loss"]) else "—")
        c5.metric("Expectancy", f"{m['expectancy']:.2f}" if pd.notna(m["expectancy"]) else "—")
        if st.button("AI Summary & Bias Diagnosis", type="primary"):
            prompt = f"""
            You are a trading performance coach.
            Data sample:
            {df.head(40).to_csv(index=False)}
            KPIs: {m}
            Provide:
            - Win/loss patterns & setups
            - Behavioral biases (confirmation, revenge, etc.)
            - 10 habit changes, 5 risk rules, 3 review routines
            - A weekly scorecard template
            """
            out = ai_complete(prompt, temperature=0.3)
            if out:
                st.markdown(out)
                sections = [("KPIs", str(m)), ("AI Diagnosis", out)]
                download_buttons("Trading_Journal_Report", sections)
            else:
                st.markdown("—")

def page_ad_copy_generator():
    hero_bar()
    header("📝 Ad Copy & Creative Generator", "PAS/AIDA + hooks and creatives")
    prod = st.text_area("Product/service description")
    audience = st.text_input("Target audience", placeholder="e.g., SMEs in Lagos")
    styles = st.multiselect("Styles", ["PAS", "AIDA", "Story-led", "Benefit bullets"], default=["PAS", "AIDA"])
    plats = st.multiselect("Platforms", ["Facebook", "Instagram", "X/Twitter", "LinkedIn", "TikTok", "YouTube"], default=["Instagram", "LinkedIn"])
    if st.button("Generate Copy & Creatives", type="primary"):
        if not prod or not audience:
            warn("Add product and audience.")
            return
        prompt = f"""
        Product: {prod}
        Audience: {audience}
        Styles: {', '.join(styles)}
        Platforms: {', '.join(plats)}
        Create:
        - 3 hooks
        - PAS ad (short + long)
        - AIDA ad (short + long)
        - 10 benefit bullets
        - Creative suggestions per platform
        - 5 CTAs
        """
        out = ai_complete(prompt, temperature=0.7)
        if out:
            st.markdown(out)
            sections = [("Product", prod), ("Audience", audience), ("Output", out)]
            download_buttons("Ad_Copy_Creative_Output", sections)
        else:
            st.markdown("—")

def page_cv_optimizer():
    hero_bar()
    header("📄 AI CV Optimizer", "ATS-friendly, quantified, tailored")
    uploaded = st.file_uploader("Upload CV (TXT/DOCX/CSV/MD)", type=["txt", "docx", "csv", "md", "markdown"])
    job_desc = st.text_area("Target job description")
    cv_text = ""
    if uploaded:
        ext = uploaded.name.split(".")[-1].lower()
        try:
            if ext == "docx":
                doc = Document(uploaded)
                cv_text = "\n".join(p.text for p in doc.paragraphs)
            elif ext in ["txt", "md", "markdown"]:
                cv_text = uploaded.read().decode("utf-8", errors="ignore")
            elif ext == "csv":
                df = pd.read_csv(uploaded)
                cv_text = df.to_csv(index=False)
            st.success("CV uploaded successfully.")
        except Exception as e:
            warn(f"Could not parse file: {e}")
    if st.button("Rewrite CV (ATS Optimized)", type="primary"):
        if not cv_text:
            warn("Please upload a CV.")
            return
        prompt = f"""
        Candidate CV:
        {cv_text}
        Target Job:
        {job_desc}
        Rewrite the CV to be ATS-friendly with quantified bullet points, strong verbs,
        and relevant keywords. Use clean markdown and clear sections.
        """
        out = ai_complete(prompt, temperature=0.25)
        if out:
            st.markdown(out)
            sections = [("Target Job", job_desc), ("Optimized CV (Markdown)", out)]
            download_buttons("AI_Optimized_CV", sections)
        else:
            st.markdown("—")

def page_job_search_copilot():
    hero_bar()
    header("💼 Job Search Copilot", "Find roles, tailor outreach, track applications")
    skills = st.text_area("Your skills (comma-separated)")
    goals = st.text_area("Career goals/preferences")
    keywords = st.text_input("Search keywords", placeholder="e.g., Data Analyst fintech remote Lagos")
    if "job_tracker" not in st.session_state:
        st.session_state.job_tracker = pd.DataFrame(columns=["Company", "Role", "Link", "Status", "Notes"])
    cA, cB = st.columns(2)
    with cA:
        company = st.text_input("Add to Tracker: Company")
        role = st.text_input("Role")
    with cB:
        link = st.text_input("Link")
        status = st.selectbox("Status", ["Planning", "Applied", "Interview", "Offer", "Rejected"])
    if st.button("Add to Tracker"):
        if company and role:
            new = pd.DataFrame([{"Company": company, "Role": role, "Link": link, "Status": status, "Notes": ""}])
            st.session_state.job_tracker = pd.concat([st.session_state.job_tracker, new], ignore_index=True)
        else:
            warn("Enter at least Company and Role.")
    st.subheader("Application Tracker")
    st.dataframe(st.session_state.job_tracker, use_container_width=True)
    if st.button("AI: Curate Roles + Cover Letter + Recruiter DM", type="primary"):
        prompt = f"""
        Skills: {skills}
        Goals/Prefs: {goals}
        Search keywords: {keywords}
        Output:
        - 10 role archetypes & where to find them
        - Boolean search strings
        - Tailored cover letter (markdown)
        - Short recruiter outreach DM
        - 2-week application plan
        """
        out = ai_complete(prompt, temperature=0.35)
        if out:
            st.markdown(out)
            sections = [("Skills", skills), ("Goals", goals), ("AI Job Pack", out)]
            download_buttons("Job_Search_Copilot_Pack", sections)
        else:
            st.markdown("—")

def page_remote_work_companion():
    hero_bar()
    header("🏡 Remote Work Companion", "Focus plan, breaks, and wellness safeguards")
    schedule = st.text_area("Work schedule", placeholder="Mon–Fri 9–5, gym Tue/Thu 6pm, caregiving Wed 3–5pm")
    calendar = st.text_area("Upcoming calendar snippets (optional)")
    sensitivity = st.select_slider("Fatigue sensitivity", ["Low", "Medium", "High"], value="Medium")
    timezone = st.text_input("Timezone (IANA)", value="Africa/Lagos")
    if st.button("Generate Focus Plan & Protocol", type="primary"):
        prompt = f"""
        Schedule: {schedule}
        Calendar notes: {calendar}
        Fatigue sensitivity: {sensitivity}
        Timezone: {timezone}
        Create:
        - Daily focus plan (time blocks)
        - Break reminders (Pomodoro + posture/hydration cues)
        - Micro-stretching routine
        - End-of-day checklist
        - Weekly productivity report template
        """
        out = ai_complete(prompt, temperature=0.25)
        if out:
            st.markdown(out)
            sections = [("Inputs", f"{schedule}\n{calendar}\nFatigue: {sensitivity}\nTZ: {timezone}"), ("AI Focus Plan", out)]
            download_buttons("Remote_Work_Companion", sections)
        else:
            st.markdown("—")

# -------------------------
# MAIN NAVIGATION
# -------------------------
def main():
    with st.sidebar:
        st.markdown("## 🧭 Navigation")
        if not st.session_state.auth["logged_in"]:
            auth_page = st.radio("Account", ["Login", "Sign Up"], label_visibility="collapsed")
            st.markdown("---")
            st.caption("Demo: Use signup to create an account.")
        else:
            st.success(f"👤 {st.session_state.auth['user']}")
            if st.button("Logout"):
                logout()
            st.markdown("---")

        pages = [
            "🛒 Marketplace",
            "🚀 Startup Copilot",
            "📊 Business Ops Monitor",
            "🎯 Work Productivity Coach",
            "💹 AI Investment Assistant",
            "📈 AI Stock Trading Journal",
            "📝 Ad Copy & Creative Generator",
            "📄 AI CV Optimizer",
            "💼 Job Search Copilot",
            "🏡 Remote Work Companion",
            "✍️ AI Content Writer",
            "📢 AI Ad Generator",
            "🎨 AI Logo Creator",
            "🔍 AI Logo Scanner",
        ]
        page = st.radio("Modules", pages, index=0, label_visibility="collapsed") if st.session_state.auth["logged_in"] else None

    if not st.session_state.auth["logged_in"]:
        if auth_page == "Login":
            login_page()
        else:
            signup_page()
        return

    try:
        page_func_map = {
            "🛒 Marketplace": page_marketplace,
            "🚀 Startup Copilot": page_startup_copilot,
            "📊 Business Ops Monitor": page_business_ops_monitor,
            "🎯 Work Productivity Coach": page_work_productivity_coach,
            "💹 AI Investment Assistant": page_investment_assistant,
            "📈 AI Stock Trading Journal": page_trading_journal,
            "📝 Ad Copy & Creative Generator": page_ad_copy_generator,
            "📄 AI CV Optimizer": page_cv_optimizer,
            "💼 Job Search Copilot": page_job_search_copilot,
            "🏡 Remote Work Companion": page_remote_work_companion,
            "✍️ AI Content Writer": page_content_writer,
            "📢 AI Ad Generator": page_ai_ad_generator,
            "🎨 AI Logo Creator": page_logo_creator,
            "🔍 AI Logo Scanner": page_logo_scanner,
        }
        if page in page_func_map:
            page_func_map[page]()
    except Exception as e:
        err(f"Something went wrong: {e}")

if __name__ == "__main__":
    main()
